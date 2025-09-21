from docx import Document

# Canonical fields for extraction
FIELDS = [
    "Counterparty",
    "Initial Valuation Date",
    "Notional",
    "Valuation Date",
    "Maturity",
    "Underlying",
    "Coupon",
    "Barrier",
    "Calendar"
]

# Aliases for variations in the document
ALIASES = {
    "Counterparty": ["Party A", "Party B"],
    "Notional": ["Notional", "Notional Amount", "Notional Amount (N)"],
    "Coupon": ["Coupon", "Coupon (C)"],
    "Barrier": ["Barrier", "Barrier (B)"],
    "Calendar": ["Business Day"],
    "Maturity": ["Termination Date", "Maturity"],
}


def parse_docx(path: str):
    """
    Parse a DOCX file and extract predefined financial fields.
    """
    doc = Document(path)
    entities = {field: None for field in FIELDS}

    # Collect all paragraph texts
    lines = [p.text for p in doc.paragraphs]

    # Include table texts in "Key ► Value" format
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            lines.append(f"{row.cells[0].text} ► {row.cells[1].text}")

    # Search for each field using its aliases
    for field in FIELDS:
        for line in lines:
            for alias in ALIASES.get(field, [field]):
                if alias.lower() in line.lower():
                    # Split line to get the value
                    if "►" in line:
                        parts = line.split("►")
                    elif "\t" in line:
                        parts = line.split("\t")
                    elif ":" in line:
                        parts = line.split(":")
                    else:
                        parts = [line, ""]  # fallback

                    value = parts[1].strip() if len(parts) > 1 else None
                    if value:
                        entities[field] = value
                        break
            if entities[field]:
                break

    return entities
